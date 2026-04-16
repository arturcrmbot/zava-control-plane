"""Populate WPP-RFP-Response-Master.docx with authored content from content/authored/."""
from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx import Document as _D

from helpers.docx_populate.content_plan import (
    SECTIONS, INLINE_PLACEHOLDERS, CELL_FILLS_BY_HEADER,
)
from helpers.docx_populate.heading_matcher import find_section
from helpers.docx_populate.md_to_docx import render_md_into_doc
from helpers.docx_populate.placeholders import resolve_inline_placeholder


REPO_ROOT = Path(__file__).parent.parent
DEFAULT_MASTER = Path(
    r"C:\Users\arzielinski\OneDrive - Microsoft\WPP Account Team - WPP_ET_AgenticAI_RFP"
    r"\MSFT_Response\WPP-RFP-Response-Master.docx"
)


def _append_md_at_section_end(doc, bounds, md_content: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    anchor_el = children[bounds.end_index] if bounds.end_index < len(children) else None

    tmp = _D()
    render_md_into_doc(tmp, md_content)
    new_elements = [p._element for p in tmp.paragraphs if p.text.strip()]

    for el in new_elements:
        if anchor_el is None:
            body.append(el)
        else:
            anchor_el.addprevious(el)


def _apply_cell_fills_by_header(doc) -> int:
    filled = 0
    for table in doc.tables:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if "Our Section(s)" not in header_cells:
            continue
        target_col = header_cells.index("Our Section(s)")
        for row in table.rows[1:]:
            first_cell_text = row.cells[0].text.strip()
            target_cell = row.cells[target_col]
            if target_cell.text.strip():
                continue
            for prefix, header_name, value in CELL_FILLS_BY_HEADER:
                if header_name != "Our Section(s)":
                    continue
                if first_cell_text.startswith(prefix) or prefix in first_cell_text:
                    target_cell.text = value
                    filled += 1
                    break
    return filled


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Populate WPP RFP master docx.")
    parser.add_argument("--master", type=Path, default=DEFAULT_MASTER,
                        help="Path to the master docx (default: OneDrive path).")
    parser.add_argument("--dry-run", action="store_true",
                        help="Show plan without writing output.")
    args = parser.parse_args(argv)

    if not args.master.exists():
        print(f"ERROR: master not found: {args.master}", file=sys.stderr)
        return 2

    doc = Document(args.master)

    plan_lines: list[str] = []
    sections_hit: list[str] = []
    sections_missed: list[str] = []
    placeholders_hit: list[str] = []
    placeholders_missed: list[str] = []

    for section in SECTIONS:
        bounds = find_section(doc, section.heading_text)
        if bounds is None:
            sections_missed.append(section.heading_text)
            plan_lines.append(f"SKIP (heading not found): {section.heading_text}")
            continue
        if not section.content_md_path.exists():
            sections_missed.append(section.heading_text)
            plan_lines.append(f"SKIP (content file missing): {section.content_md_path.name}")
            continue
        preview = section.content_md_path.read_text(encoding="utf-8")[:80].replace("\n", " ")
        plan_lines.append(f"INJECT: {section.heading_text}  ->  {preview}...")
        sections_hit.append(section.heading_text)

    for rule in INLINE_PLACEHOLDERS:
        found = any(p.text.strip().startswith(rule.match_prefix.strip()) for p in doc.paragraphs)
        if found:
            placeholders_hit.append(rule.match_prefix)
            plan_lines.append(f"{rule.resolution.upper()}: {rule.match_prefix[:60]}...")
        else:
            placeholders_missed.append(rule.match_prefix)
            plan_lines.append(f"SKIP (placeholder not found): {rule.match_prefix[:60]}...")

    print("=== POPULATION PLAN ===")
    for line in plan_lines:
        print("  " + line)
    print(f"Sections to inject: {len(sections_hit)}  /  Missed: {len(sections_missed)}")
    print(f"Placeholders to resolve: {len(placeholders_hit)}  /  Missed: {len(placeholders_missed)}")

    if args.dry_run:
        print("Dry-run: no output file written.")
        return 0 if not sections_missed and not placeholders_missed else 1

    # Execute — re-open master fresh
    doc = Document(args.master)

    for rule in INLINE_PLACEHOLDERS:
        resolve_inline_placeholder(doc, rule)

    cells_filled = _apply_cell_fills_by_header(doc)

    for section in SECTIONS:
        bounds = find_section(doc, section.heading_text)
        if bounds is None or not section.content_md_path.exists():
            continue
        md_content = section.content_md_path.read_text(encoding="utf-8")
        _append_md_at_section_end(doc, bounds, md_content)

    timestamp = datetime.now().strftime("%Y-%m-%d-%H%M")
    output = args.master.parent / f"WPP-RFP-Response-Master-populated-{timestamp}.docx"
    doc.save(output)

    report_path = output.with_suffix(".report.md")
    with report_path.open("w", encoding="utf-8") as fh:
        fh.write(f"# Populate Report — {timestamp}\n\n")
        fh.write(f"## Output\n{output}\n\n")
        fh.write(f"## Sections injected ({len(sections_hit)})\n")
        for h in sections_hit:
            fh.write(f"- {h}\n")
        fh.write(f"\n## Sections missed ({len(sections_missed)})\n")
        for h in sections_missed:
            fh.write(f"- {h}\n")
        fh.write(f"\n## Placeholders resolved ({len(placeholders_hit)})\n")
        for h in placeholders_hit:
            fh.write(f"- {h[:80]}\n")
        fh.write(f"\n## Placeholders missed ({len(placeholders_missed)})\n")
        for h in placeholders_missed:
            fh.write(f"- {h[:80]}\n")
        fh.write(f"\n## Table cells filled\n{cells_filled}\n")

    print(f"\nOutput: {output}")
    print(f"Report: {report_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
