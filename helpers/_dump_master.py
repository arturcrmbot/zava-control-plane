"""Dump WPP master docx content for triage."""
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

p = r"C:\Users\arzielinski\OneDrive - Microsoft\WPP Account Team - WPP_ET_AgenticAI_RFP\MSFT_Response\WPP-RFP-Response-Master.docx"
out = Path(r"c:\dev\ghcp sdk stuff\scratch\wpp-master-extract.md")
out.parent.mkdir(exist_ok=True)

d = Document(p)
lines = []

# index paragraphs and tables together in document order
# docx has tables as separate element tree — interleave via XML order
from docx.oxml.ns import qn
body = d.element.body

def iter_block_items(parent):
    for child in parent.iterchildren():
        tag = child.tag
        if tag == qn('w:p'):
            yield ('p', child)
        elif tag == qn('w:tbl'):
            yield ('t', child)

para_map = {p._element: p for p in d.paragraphs}
tbl_map = {t._element: t for t in d.tables}

for kind, el in iter_block_items(body):
    if kind == 'p':
        para = para_map.get(el)
        if not para:
            continue
        s = para.style.name
        t = para.text.strip()
        if not t:
            continue
        if s == "heading 10":
            lines.append(f"\n\n{'='*80}\n## {t}\n{'='*80}")
        elif s == "heading 20":
            lines.append(f"\n### {t}\n{'-'*60}")
        elif s == "heading 30":
            lines.append(f"\n#### {t}")
        elif s == "List Paragraph":
            lines.append(f"  - {t}")
        else:
            lines.append(t)
    elif kind == 't':
        tbl = tbl_map.get(el)
        if not tbl:
            continue
        lines.append("\n[TABLE]")
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
            lines.append("  | " + " | ".join(cells))
        lines.append("[/TABLE]\n")

out.write_text("\n".join(lines), encoding='utf-8')
print(f"{out} -> {len(lines)} lines, {out.stat().st_size} bytes")
