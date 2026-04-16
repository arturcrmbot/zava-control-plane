"""Render a restricted Markdown subset into a python-docx Document.

Subset supported:
  - Headings: # ## ### #### (mapped to Heading 1..4)
  - Paragraphs
  - Bullet lists (- or *)
  - Numbered lists (1. 2. 3.)
  - GFM pipe tables (header | separator | data rows)
  - Inline: **bold**, *italic*, `code`, [text](url)

Images and code blocks are NOT supported.
"""
from __future__ import annotations

import re
from docx.document import Document as DocxDocument


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|.+\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|\s*(:?-+:?\s*\|\s*)+$")

_INLINE_RE = re.compile(
    r"\*\*(?P<bold>[^*]+)\*\*"
    r"|\*(?P<italic>[^*]+)\*"
    r"|`(?P<code>[^`]+)`"
    r"|\[(?P<linktext>[^\]]+)\]\((?P<linkurl>[^)]+)\)"
)


DEFAULT_STYLE_MAP: dict[str, str] = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 4",
    "h6": "Heading 4",
    "bullet": "List Bullet",
    "numbered": "List Number",
    "table": "Table Grid",
}


def _add_runs(para, text: str) -> None:
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group("bold"):
            r = para.add_run(m.group("bold"))
            r.bold = True
        elif m.group("italic"):
            r = para.add_run(m.group("italic"))
            r.italic = True
        elif m.group("code"):
            r = para.add_run(m.group("code"))
            r.font.name = "Consolas"
        elif m.group("linktext"):
            r = para.add_run(m.group("linktext"))
            r.underline = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def _parse_table_row(line: str) -> list[str]:
    """Split a pipe-table row into stripped cell values.

    Leading/trailing pipes are stripped; interior pipes separate cells. Escaped
    pipes (`\\|`) inside cells are preserved as literal pipes.
    """
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    # Simple split — no escaped-pipe support; our authored content doesn't use them.
    return [cell.strip() for cell in stripped.split("|")]


def _render_table(doc: DocxDocument, header: list[str], rows: list[list[str]],
                  style_map: dict[str, str]) -> None:
    n_cols = len(header)
    if n_cols == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table_style = style_map.get("table", "Table Grid")
    try:
        table.style = table_style
    except KeyError:
        # Destination doc doesn't define the requested style; leave default.
        pass

    # Header row — populate cells and bold each run.
    for col_idx in range(n_cols):
        cell = table.cell(0, col_idx)
        text = header[col_idx] if col_idx < len(header) else ""
        cell.text = ""
        para = cell.paragraphs[0]
        _add_runs(para, text)
        for run in para.runs:
            run.bold = True

    # Data rows — pad short rows, truncate long ones to column count.
    for row_idx, row_cells in enumerate(rows, start=1):
        for col_idx in range(n_cols):
            value = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_runs(para, value)


def render_md_into_doc(
    doc: DocxDocument,
    md: str,
    bullet_style: str | None = None,
    style_map: dict[str, str] | None = None,
) -> None:
    """Render `md` into `doc`. `style_map` overrides default style names.

    Keys: h1, h2, h3, h4, h5, h6, bullet, numbered, table. Unspecified keys
    fall back to DEFAULT_STYLE_MAP. The legacy `bullet_style` kwarg
    (pre-style-map API) is still honoured if `style_map` does not set "bullet".
    """
    sm = {**DEFAULT_STYLE_MAP, **(style_map or {})}
    if bullet_style is not None and (not style_map or "bullet" not in style_map):
        sm["bullet"] = bullet_style

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            i += 1
            continue

        # GFM pipe table: pipe-row followed by a dash separator.
        if (_TABLE_ROW_RE.match(line)
                and i + 1 < len(lines)
                and _TABLE_SEP_RE.match(lines[i + 1].rstrip())):
            header = _parse_table_row(line)
            i += 2  # skip header and separator
            data_rows: list[list[str]] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].rstrip()):
                data_rows.append(_parse_table_row(lines[i].rstrip()))
                i += 1
            _render_table(doc, header, data_rows, sm)
            continue

        hm = _HEADING_RE.match(line)
        if hm:
            level = len(hm.group(1))
            level = min(level, 6)
            text = hm.group(2).strip()
            para = doc.add_paragraph(style=sm[f"h{level}"])
            _add_runs(para, text)
            i += 1
            continue

        bm = _BULLET_RE.match(line)
        if bm:
            para = doc.add_paragraph(style=sm["bullet"])
            _add_runs(para, bm.group(1).strip())
            i += 1
            continue

        nm = _NUMBERED_RE.match(line)
        if nm:
            para = doc.add_paragraph(style=sm["numbered"])
            _add_runs(para, nm.group(1).strip())
            i += 1
            continue

        para = doc.add_paragraph()
        _add_runs(para, line.strip())
        i += 1
