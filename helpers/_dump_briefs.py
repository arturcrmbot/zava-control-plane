"""Dump all four WPP brief docx files for digest creation."""
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

BASE = r"C:\Users\arzielinski\OneDrive - Microsoft\WPP Account Team - WPP_ET_AgenticAI_RFP\RFx Documentation - Originals"
DOCS = {
    "vendor-brief": rf"{BASE}\WPPET-1-Enterprise-Agent-Framework-Vendor-Brief-260331.docx",
    "poc1-finance": rf"{BASE}\WPPET-POC1-Finance-Expense-Compliance-260331.docx",
    "poc2-hr": rf"{BASE}\WPPET-POC2-HR-Workforce-Transformation-260331.docx",
    "qa-clarification": rf"{BASE}\WPPET-5-Enterprise-Agent-Framework-QA-260408.docx",
}

outdir = Path(r"c:\dev\ghcp sdk stuff\scratch")
outdir.mkdir(exist_ok=True)


def dump(path, label):
    d = Document(path)
    lines = []
    body = d.element.body
    para_map = {p._element: p for p in d.paragraphs}
    tbl_map = {t._element: t for t in d.tables}

    for child in body.iterchildren():
        tag = child.tag
        if tag == qn('w:p'):
            para = para_map.get(child)
            if not para:
                continue
            s = para.style.name if para.style else ""
            t = para.text.strip()
            if not t:
                continue
            if s.lower().startswith("heading 1"):
                lines.append(f"\n\n{'='*80}\n# {t}\n{'='*80}")
            elif s.lower().startswith("heading 2"):
                lines.append(f"\n## {t}\n{'-'*60}")
            elif s.lower().startswith("heading 3"):
                lines.append(f"\n### {t}")
            elif s.lower().startswith("heading 4"):
                lines.append(f"\n#### {t}")
            elif s == "List Paragraph":
                lines.append(f"  - {t}")
            else:
                lines.append(f"[{s}] {t}" if s and s != "Normal" else t)
        elif tag == qn('w:tbl'):
            tbl = tbl_map.get(child)
            if not tbl:
                continue
            lines.append("\n[TABLE]")
            for row in tbl.rows:
                cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
                lines.append("  | " + " | ".join(cells))
            lines.append("[/TABLE]\n")

    outpath = outdir / f"brief-raw-{label}.md"
    outpath.write_text("\n".join(lines), encoding='utf-8')
    print(f"{label}: {outpath} -> {len(lines)} lines, {outpath.stat().st_size} bytes")


for label, path in DOCS.items():
    try:
        dump(path, label)
    except Exception as e:
        print(f"ERROR {label}: {e}")
