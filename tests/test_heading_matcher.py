from docx import Document
from helpers.docx_populate.heading_matcher import find_section, SectionBounds


def test_find_existing_section_returns_bounds(fixtures_dir):
    doc = Document(fixtures_dir / "mini-master.docx")
    bounds = find_section(doc, "Section A")
    assert bounds is not None
    assert bounds.heading_index >= 0
    assert bounds.end_index > bounds.heading_index


def test_find_missing_section_returns_none(fixtures_dir):
    doc = Document(fixtures_dir / "mini-master.docx")
    assert find_section(doc, "Nonexistent") is None


def test_section_bounds_end_points_to_document_end_for_last_section(fixtures_dir):
    doc = Document(fixtures_dir / "mini-master.docx")
    bounds = find_section(doc, "Section C (empty)")
    assert bounds is not None
    total_blocks = sum(1 for _ in doc.element.body.iterchildren())
    assert bounds.end_index == total_blocks
