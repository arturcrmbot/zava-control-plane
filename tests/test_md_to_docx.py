from docx import Document
from helpers.docx_populate.md_to_docx import render_md_into_doc


def test_render_paragraphs_and_headings(tmp_path):
    doc = Document()
    md = """## Heading Two

A paragraph with **bold** and *italic*.

### Heading Three

- Bullet one
- Bullet two

Another paragraph.
"""
    render_md_into_doc(doc, md)
    out = tmp_path / "out.docx"
    doc.save(out)
    reloaded = Document(out)
    styles = [(p.style.name, p.text) for p in reloaded.paragraphs if p.text.strip()]
    assert styles[0][0].lower().startswith("heading")
    assert styles[0][1] == "Heading Two"
    assert any(p.text == "Bullet one" for p in reloaded.paragraphs)


def test_render_preserves_bold(tmp_path):
    doc = Document()
    md = "A **bold** word."
    render_md_into_doc(doc, md)
    para = doc.paragraphs[-1]
    runs = list(para.runs)
    assert any(r.bold and r.text == "bold" for r in runs)


def test_render_gfm_table(tmp_path):
    doc = Document()
    md = """## Table test

| Header A | Header B | Header C |
|---|---|---|
| row 1 a | row 1 b | row 1 c |
| row 2 a | row 2 b | row 2 c |

Trailing paragraph.
"""
    render_md_into_doc(doc, md)
    out = tmp_path / "out.docx"
    doc.save(out)
    reloaded = Document(out)

    # Exactly one table with 3 rows × 3 cols.
    assert len(reloaded.tables) == 1
    table = reloaded.tables[0]
    assert len(table.rows) == 3
    assert len(table.columns) == 3

    # Header cells populated and bolded.
    assert table.cell(0, 0).text == "Header A"
    assert table.cell(0, 2).text == "Header C"
    header_runs = table.cell(0, 0).paragraphs[0].runs
    assert any(r.bold for r in header_runs)

    # Data cells populated.
    assert table.cell(1, 1).text == "row 1 b"
    assert table.cell(2, 2).text == "row 2 c"

    # Trailing paragraph rendered after the table.
    texts = [p.text for p in reloaded.paragraphs]
    assert any(t == "Trailing paragraph." for t in texts)


def test_render_uses_custom_style_map(tmp_path):
    """A caller-supplied style_map overrides the default heading style names."""
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    # The target style must exist in the document's style catalogue before
    # python-docx can apply it. Real masters define their own styles; for this
    # test we register one explicitly.
    doc.styles.add_style("heading 20", WD_STYLE_TYPE.PARAGRAPH)

    md = "## Heading\n\nBody"
    render_md_into_doc(doc, md, style_map={"h2": "heading 20"})

    out = tmp_path / "out.docx"
    doc.save(out)
    reloaded = Document(out)

    styled = [(p.style.name, p.text) for p in reloaded.paragraphs if p.text.strip()]
    heading_entry = next(s for s in styled if s[1] == "Heading")
    assert heading_entry[0] == "heading 20"
