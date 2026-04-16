from docx import Document
from helpers.docx_populate.placeholders import resolve_inline_placeholder, PlaceholderRule


def test_replace_placeholder_paragraph(tmp_path, fixtures_dir):
    doc = Document(fixtures_dir / "mini-master.docx")
    rule = PlaceholderRule(
        match_prefix="<<Fill me please>>",
        resolution="replace_with",
        content="This is the replacement content.",
    )
    found = resolve_inline_placeholder(doc, rule)
    assert found >= 1
    out = tmp_path / "out.docx"
    doc.save(out)
    reloaded = Document(out)
    texts = [p.text for p in reloaded.paragraphs]
    assert "<<Fill me please>>" not in " ".join(texts)
    assert any("This is the replacement content." in t for t in texts)


def test_delete_placeholder_paragraph(tmp_path, fixtures_dir):
    doc = Document(fixtures_dir / "mini-master.docx")
    rule = PlaceholderRule(
        match_prefix="Cut and paste below",
        resolution="delete_paragraph",
    )
    found = resolve_inline_placeholder(doc, rule)
    assert found >= 1
    out = tmp_path / "out.docx"
    doc.save(out)
    reloaded = Document(out)
    texts = [p.text for p in reloaded.paragraphs]
    assert not any("Cut and paste" in t for t in texts)


def test_missing_placeholder_returns_zero(fixtures_dir):
    doc = Document(fixtures_dir / "mini-master.docx")
    rule = PlaceholderRule(
        match_prefix="<<Not actually in the doc>>",
        resolution="replace_with",
        content="irrelevant",
    )
    assert resolve_inline_placeholder(doc, rule) == 0


def test_mid_paragraph_replace_preserves_surrounding_text(tmp_path):
    """A marker appearing inside a paragraph should leave surrounding text
    intact and insert the rendered replacement AFTER the modified paragraph."""
    doc = Document()
    doc.add_paragraph("Before text <<INLINE>> after text.")

    rule = PlaceholderRule(
        match_prefix="<<INLINE>>",
        resolution="replace_with",
        content="## Injected heading\n\nReplacement body.",
    )
    count = resolve_inline_placeholder(doc, rule)
    assert count == 1

    out = tmp_path / "out.docx"
    doc.save(out)
    reloaded = Document(out)

    texts = [p.text for p in reloaded.paragraphs if p.text.strip()]
    # Original paragraph present with marker removed, surrounding text preserved.
    assert any("Before text" in t and "after text" in t and "<<INLINE>>" not in t
               for t in texts), f"expected cleaned original paragraph in {texts}"

    # The injected heading and body should appear after the original paragraph.
    original_idx = next(i for i, t in enumerate(texts)
                        if "Before text" in t and "after text" in t)
    assert texts[original_idx + 1] == "Injected heading"
    assert texts[original_idx + 2] == "Replacement body."

    # Heading paragraph must actually use a heading style.
    styled = [(p.style.name, p.text) for p in reloaded.paragraphs if p.text.strip()]
    heading_entry = next(s for s in styled if s[1] == "Injected heading")
    assert heading_entry[0].lower().startswith("heading")
